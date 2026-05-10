"""report_exporter.py

Generates a formal Word (.docx) evaluation report from a run_recommendation() result.

Sections:
  1. Cover
  2. Student profile
  3. Six-dimension capability portrait (table + radar chart image)
  4. Path fitness analysis (table + bar chart image)
  5. Main & alternative paths
  6. Recommendation rationale (algorithm + intent analysis)
  7. 1-3 year planning suggestions
  8. Similar senior cases
  9. Employment / internship / postgrad recommendations
  10. Disclaimer
"""
from __future__ import annotations

import io
import math
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# matplotlib is optional — charts are skipped gracefully if unavailable
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    _HAS_MPL = True
except ImportError:
    _HAS_MPL = False

from .config import DIMENSIONS, DIMENSION_LABELS

# ── constants ──────────────────────────────────────────────────────────────

SYSTEM_NAME = "AI辅导员智能体系统"
DISCLAIMER = (
    "本报告由 AI辅导员智能体系统基于学生提供信息、历史案例、政策规则和算法模型辅助生成，"
    "仅作为成长规划参考。涉及推免、专项计划、就业派遣等正式事项，"
    "请以学校和学院最新政策通知及辅导员审核意见为准。"
)

DIM_LABELS_FULL = {
    "academic_foundation": "学业基础能力",
    "research_innovation": "科研创新能力",
    "competition_practice": "竞赛实践能力",
    "engineering_practice": "工程实践与实习能力",
    "organization_service": "组织服务与学生工作能力",
    "growth_planning": "成长潜力与规划准备度",
}

PATH_ORDER = [
    "普通推免", "特殊专长A类", "特殊专长B/本硕博",
    "特殊专长C/辅导员计划", "支教计划", "工程专项/专硕",
    "实习就业", "考研", "考公",
]


# ── helpers ────────────────────────────────────────────────────────────────

def _safe(v: Any, default: str = "暂无") -> str:
    if v is None:
        return default
    s = str(v).strip()
    return s if s and s not in ("None", "[]", "{}", "null") else default


def _list_to_str(v: Any, default: str = "暂无") -> str:
    if not v:
        return default
    if isinstance(v, list):
        items = [str(x).strip() for x in v if x and str(x).strip()]
        return "；".join(items) if items else default
    return _safe(v, default)


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _para(doc: Document, text: str, bold: bool = False, size: int = 11,
          color: str | None = None, indent: float = 0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def _table_header_row(table, headers: list[str], bg: str = "1d4ed8"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_bg(cell, bg)


# ── chart generators ───────────────────────────────────────────────────────

def _radar_chart_image(scores: dict) -> io.BytesIO | None:
    if not _HAS_MPL:
        return None
    labels = [DIM_LABELS_FULL.get(d, d) for d in DIMENSIONS]
    values = [float(scores.get(d, 0) or 0) for d in DIMENSIONS]
    N = len(labels)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    angles += angles[:1]
    values_plot = values + values[:1]

    fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
    ax.set_theta_offset(math.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    font = _get_font()
    tick_kw = {"fontsize": 8}
    if font:
        tick_kw["fontproperties"] = font
    ax.set_xticklabels(labels, **tick_kw)
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(["20", "40", "60", "80", "100"], fontsize=7)
    ax.plot(angles, values_plot, "o-", linewidth=2, color="#1d4ed8")
    ax.fill(angles, values_plot, alpha=0.25, color="#1d4ed8")
    title_kw = {"size": 11, "pad": 15}
    if font:
        title_kw["fontproperties"] = font
    ax.set_title("六维能力画像", **title_kw)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _bar_chart_image(path_scores: dict) -> io.BytesIO | None:
    if not _HAS_MPL:
        return None
    paths = [p for p in PATH_ORDER if p in path_scores]
    scores = [path_scores[p] for p in paths]
    colors = ["#1d4ed8" if s == max(scores) else "#93c5fd" for s in scores]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.barh(paths, scores, color=colors)
    ax.set_xlim(0, 105)
    font = _get_font()
    xlabel_kw = {}
    if font:
        xlabel_kw["fontproperties"] = font
    ax.set_xlabel("路径适配度 / 100", **xlabel_kw)
    title_kw = {}
    if font:
        title_kw["fontproperties"] = font
    ax.set_title("各路径适配度", **title_kw)
    for bar, score in zip(bars, scores):
        ax.text(score + 1, bar.get_y() + bar.get_height() / 2,
                f"{score:.1f}", va="center", fontsize=8)
    if font:
        for label in ax.get_yticklabels():
            label.set_fontproperties(font)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _get_font():
    """Return a matplotlib FontProperties for CJK text, falling back gracefully.

    Tries a prioritised list of CJK fonts available on Windows, macOS, and
    common Linux distributions.  If none is found, returns None so matplotlib
    uses its default font (charts will render but CJK text may show as boxes
    on systems without any CJK font installed).
    """
    try:
        from matplotlib.font_manager import FontProperties, fontManager
        candidates = [
            "Microsoft YaHei",   # Windows
            "SimHei",            # Windows fallback
            "PingFang SC",       # macOS
            "Noto Sans CJK SC",  # Linux (fonts-noto-cjk)
            "Noto Sans SC",
            "WenQuanYi Micro Hei",  # Linux (fonts-wqy-microhei)
            "WenQuanYi Zen Hei",
            "Arial Unicode MS",
            "DejaVu Sans",       # last resort — no CJK but won't crash
        ]
        available = {f.name for f in fontManager.ttflist}
        for name in candidates:
            if name in available:
                return FontProperties(family=name)
        return None
    except Exception:
        return None


# ── section builders ───────────────────────────────────────────────────────

def _section_cover(doc: Document, generated_at: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("学生成长路径评估报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph()
    for line, size in [
        (f"生成时间：{generated_at}", 11),
        (f"系统名称：{SYSTEM_NAME}", 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()


def _section_profile(doc: Document, profile: dict):
    _heading(doc, "一、学生基础画像", level=1)

    # Support both new structured student format and old flat profile
    is_structured = "gpa_by_semester" in profile

    if is_structured:
        # Compute average GPA from semester data
        gpa_vals = [v for v in profile.get("gpa_by_semester", {}).values() if v is not None]
        gpa_avg = f"{sum(gpa_vals)/len(gpa_vals):.2f}" if gpa_vals else "暂无"
        gpa_trend = "、".join(
            f"{sem}: {v}" for sem, v in profile.get("gpa_by_semester", {}).items() if v is not None
        )
        flat = profile.get("profile") or {}
        fields = [
            ("班级", _safe(profile.get("class_name"))),
            ("学号", _safe(profile.get("student_no"))),
            ("特长", _safe(profile.get("specialty"))),
            ("去向", _safe(profile.get("destination"))),
            ("去向单位", _safe(profile.get("destination_unit"))),
            ("平均学分绩点（均值）", gpa_avg),
            ("各学期绩点", gpa_trend or "暂无"),
            ("CET-4", _safe(flat.get("cet4"))),
            ("CET-6", _safe(flat.get("cet6"))),
            ("竞赛/奖项", _list_to_str(flat.get("competitions"))),
            ("学生工作", _list_to_str(flat.get("student_work"))),
            ("志愿时长", f"{flat.get('volunteer_hours') or 0} 小时"),
        ]
    else:
        fields = [
            ("专业", _safe(profile.get("major"))),
            ("绩点", _safe(profile.get("gpa"))),
            ("CET-4", _safe(profile.get("cet4"))),
            ("CET-6", _safe(profile.get("cet6"))),
            ("竞赛经历", _list_to_str(profile.get("competitions"))),
            ("科研/论文/专利", _list_to_str(profile.get("papers")) + "；" + _list_to_str(profile.get("research_experiences")) if (profile.get("papers") or profile.get("research_experiences")) else "暂无"),
            ("实习/项目经历", _list_to_str(profile.get("internships")) + "；" + _list_to_str(profile.get("projects")) if (profile.get("internships") or profile.get("projects")) else "暂无"),
            ("学生工作/志愿服务", _list_to_str(profile.get("student_work")) + f"（志愿时长 {profile.get('volunteer_hours') or 0} 小时）"),
            ("目标意向", _safe(profile.get("target"))),
        ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_bg(row.cells[0], "EFF6FF")
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph()

    # For structured students: add semester detail tables
    if is_structured:
        _section_semester_details(doc, profile)


def _section_semester_details(doc: Document, student: dict):
    """Render per-semester GPA, positions, awards, and evaluations tables."""
    semesters = ["大一上", "大一下", "大二上", "大二下", "大三上", "大三下", "大四上"]

    _heading(doc, "各学期详细数据", level=2)

    # ── GPA trend table ──────────────────────────────────────────────────
    gpa_map = student.get("gpa_by_semester") or {}
    _para(doc, "平均学分绩点（各学期）", bold=True)
    gpa_table = doc.add_table(rows=2, cols=len(semesters))
    gpa_table.style = "Table Grid"
    for i, sem in enumerate(semesters):
        gpa_table.rows[0].cells[i].text = sem
        gpa_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        gpa_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(gpa_table.rows[0].cells[i], "1d4ed8")
        gpa_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        val = gpa_map.get(sem)
        gpa_table.rows[1].cells[i].text = f"{val:.2f}" if val is not None else "—"
        gpa_table.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        if val is not None and val >= 3.5:
            _set_cell_bg(gpa_table.rows[1].cells[i], "ECFDF5")
        elif val is not None and val < 2.8:
            _set_cell_bg(gpa_table.rows[1].cells[i], "FEF2F2")
    doc.add_paragraph()

    # ── Positions table ──────────────────────────────────────────────────
    pos_map = student.get("positions_by_semester") or {}
    _para(doc, "担任职务（各学期）", bold=True)
    pos_table = doc.add_table(rows=len(semesters) + 1, cols=2)
    pos_table.style = "Table Grid"
    _table_header_row(pos_table, ["学期", "担任职务"])
    for i, sem in enumerate(semesters):
        pos_table.rows[i + 1].cells[0].text = sem
        pos_table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        pos_table.rows[i + 1].cells[0].width = Cm(2.5)
        val = pos_map.get(sem) or "—"
        pos_table.rows[i + 1].cells[1].text = val
        pos_table.rows[i + 1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── Awards table ─────────────────────────────────────────────────────
    awd_map = student.get("awards_by_semester") or {}
    _para(doc, "所获奖项（各学期）", bold=True)
    awd_table = doc.add_table(rows=len(semesters) + 1, cols=2)
    awd_table.style = "Table Grid"
    _table_header_row(awd_table, ["学期", "所获奖项"])
    for i, sem in enumerate(semesters):
        awd_table.rows[i + 1].cells[0].text = sem
        awd_table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        awd_table.rows[i + 1].cells[0].width = Cm(2.5)
        val = awd_map.get(sem) or "—"
        awd_table.rows[i + 1].cells[1].text = val
        awd_table.rows[i + 1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── Evaluations table ────────────────────────────────────────────────
    evl_map = student.get("evaluations_by_semester") or {}
    _para(doc, "学期评价（各学期）", bold=True)
    evl_table = doc.add_table(rows=len(semesters) + 1, cols=2)
    evl_table.style = "Table Grid"
    _table_header_row(evl_table, ["学期", "学期评价"])
    for i, sem in enumerate(semesters):
        evl_table.rows[i + 1].cells[0].text = sem
        evl_table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        evl_table.rows[i + 1].cells[0].width = Cm(2.5)
        val = evl_map.get(sem) or "—"
        evl_table.rows[i + 1].cells[1].text = val
        evl_table.rows[i + 1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _section_six_dim(doc: Document, scores: dict):
    _heading(doc, "二、六维能力画像", level=1)

    # Table
    table = doc.add_table(rows=len(DIMENSIONS) + 1, cols=3)
    table.style = "Table Grid"
    _table_header_row(table, ["维度", "英文键", "得分（0-100）"])
    for i, dim in enumerate(DIMENSIONS):
        row = table.rows[i + 1]
        row.cells[0].text = DIM_LABELS_FULL.get(dim, dim)
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = dim
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        val = float(scores.get(dim, 0) or 0)
        row.cells[2].text = f"{val:.1f}"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        if val >= 65:
            _set_cell_bg(row.cells[2], "ECFDF5")
        elif val <= 35:
            _set_cell_bg(row.cells[2], "FEF2F2")
    doc.add_paragraph()

    # Radar chart
    img_buf = _radar_chart_image(scores)
    if img_buf:
        doc.add_picture(img_buf, width=Cm(10))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _section_path_scores(doc: Document, path_scores: dict):
    _heading(doc, "三、路径适配度分析", level=1)

    ranked = sorted(path_scores.items(), key=lambda x: x[1], reverse=True)
    table = doc.add_table(rows=len(ranked) + 1, cols=3)
    table.style = "Table Grid"
    _table_header_row(table, ["路径", "适配度（/100）", "排名"])
    for rank, (path, score) in enumerate(ranked, 1):
        row = table.rows[rank]
        row.cells[0].text = path
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = f"{score:.1f}"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = f"第 {rank} 位"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        if rank == 1:
            _set_cell_bg(row.cells[0], "DBEAFE")
            _set_cell_bg(row.cells[1], "DBEAFE")
            _set_cell_bg(row.cells[2], "DBEAFE")
    doc.add_paragraph()

    # Bar chart
    img_buf = _bar_chart_image(path_scores)
    if img_buf:
        doc.add_picture(img_buf, width=Cm(14))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _section_main_path(doc: Document, path_scores: dict, algo: dict):
    _heading(doc, "四、主推荐路径与备选路径", level=1)

    ranked = sorted(path_scores.items(), key=lambda x: x[1], reverse=True)
    main = ranked[0] if ranked else ("暂无", 0)
    alts = ranked[1:3]
    weak = ranked[-2:] if len(ranked) >= 4 else []

    _para(doc, f"主推荐路径：{main[0]}（适配度 {main[1]:.1f}/100）", bold=True, size=12)

    reasons = algo.get("top_path_reasons") or []
    if reasons:
        _para(doc, "推荐理由：", bold=True)
        for r in reasons:
            _para(doc, f"• {r}", indent=0.5)

    if alts:
        doc.add_paragraph()
        _para(doc, "备选路径：", bold=True)
        for k, v in alts:
            _para(doc, f"• {k}（{v:.1f}/100）", indent=0.5)

    if weak:
        doc.add_paragraph()
        _para(doc, "不优先推荐路径：", bold=True)
        for k, v in weak:
            _para(doc, f"• {k}（{v:.1f}/100）", indent=0.5)

    doc.add_paragraph()


def _section_rationale(doc: Document, algo: dict, intent: dict):
    _heading(doc, "五、推荐依据说明", level=1)

    # Methods
    methods = algo.get("method_used") or []
    if methods:
        _para(doc, "评估方法：" + " · ".join(methods), size=10, color="64748b")
        doc.add_paragraph()

    # Intent analysis
    detected = intent.get("detected_intents") or []
    bonus = intent.get("applied_bonus") or {}
    intent_exp = intent.get("explanation") or ""
    if detected or intent_exp:
        _para(doc, "目标意向识别", bold=True)
        if detected:
            _para(doc, "识别到的意向：" + "、".join(detected), indent=0.5)
        if bonus:
            bonus_str = "；".join(f"{p} +{v:.0f}分" for p, v in bonus.items())
            _para(doc, f"意向修正（最多 8 分）：{bonus_str}", indent=0.5)
        if intent_exp:
            _para(doc, intent_exp, indent=0.5, size=10, color="475569")
        doc.add_paragraph()

    # Strengths / weaknesses
    strengths = algo.get("strengths") or []
    weaknesses = algo.get("weaknesses") or []
    col_data = []
    if strengths:
        col_data.append(("优势维度", "、".join(strengths), "ECFDF5"))
    if weaknesses:
        col_data.append(("短板维度", "、".join(weaknesses), "FEF2F2"))
    for label, content, bg in col_data:
        _para(doc, label, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(content)
        run.font.size = Pt(10)

    # Risk factors
    risks = algo.get("risk_factors") or []
    if risks:
        doc.add_paragraph()
        _para(doc, "风险提示", bold=True)
        for r in risks:
            _para(doc, f"• {r}", indent=0.5)

    # Improvement suggestions
    suggestions = algo.get("improvement_suggestions") or []
    if suggestions:
        doc.add_paragraph()
        _para(doc, "提升建议", bold=True)
        for i, s in enumerate(suggestions, 1):
            _para(doc, f"{i}. {s}", indent=0.5)

    doc.add_paragraph()


def _section_planning(doc: Document, algo: dict, advice: str):
    _heading(doc, "六、未来 1-3 年规划建议", level=1)

    suggestions = algo.get("improvement_suggestions") or []
    total = len(suggestions)

    _para(doc, "近期（1-3 个月）", bold=True)
    near = suggestions[:2] if total >= 2 else suggestions
    if near:
        for s in near:
            _para(doc, f"• {s}", indent=0.5)
    else:
        _para(doc, "• 核对系统抽取的绩点、竞赛、论文、实习和学生工作信息，补全缺失字段。", indent=0.5)
    _para(doc, "• 与相似学长案例对照，优先提升六维能力画像中差距最大的两个维度。", indent=0.5)

    doc.add_paragraph()
    _para(doc, "中期（6-12 个月）", bold=True)
    mid = suggestions[2:4] if total >= 4 else []
    if mid:
        for s in mid:
            _para(doc, f"• {s}", indent=0.5)
    else:
        _para(doc, "• 围绕主推荐路径深化核心优势，积累可量化的成果（竞赛奖项、论文、实习证明等）。", indent=0.5)
        _para(doc, "• 专业方向作为辅助参考，结合目标路径确认技术栈或研究方向的匹配度。", indent=0.5)

    doc.add_paragraph()
    _para(doc, "长期（1-3 年）", bold=True)
    _para(doc, "• 按主推荐路径完成关键节点（如保研申请、实习转正、考研备考、考公报名等）。", indent=0.5)
    _para(doc, "• 持续关注学院最新政策通知，及时调整规划方向。", indent=0.5)

    doc.add_paragraph()


def _section_seniors(doc: Document, seniors: list[dict]):
    _heading(doc, "七、相似学长学姐案例", level=1)
    _para(doc, "以下案例已脱敏处理，隐去姓名、学号和联系方式，保留成长路径与匹配原因。",
          size=10, color="64748b")
    doc.add_paragraph()

    if not seniors:
        _para(doc, "当前历史学生库较小，暂未找到足够相似案例。")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=len(seniors[:5]) + 1, cols=5)
    table.style = "Table Grid"
    _table_header_row(table, ["案例编号", "相似度", "成长路径", "最终去向", "匹配原因"])
    for i, s in enumerate(seniors[:5]):
        row = table.rows[i + 1]
        row.cells[0].text = _safe(s.get("case_id"), f"匿名案例{i+1}")
        row.cells[1].text = f"{s.get('similarity', 0):.1f}%"
        row.cells[2].text = _safe(s.get("path_type") or s.get("growth_destination_type"))
        row.cells[3].text = _safe(s.get("destination"))
        row.cells[4].text = _safe(s.get("match_reason"))
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _section_companies(doc: Document, companies: list[dict]):
    _heading(doc, "八、就业/实习/升学建议", level=1)

    if not companies:
        _para(doc, "暂未导入校企合作单位，建议管理员上传合作单位表。")
        doc.add_paragraph()
        return

    for c in companies[:5]:
        name = _safe(c.get("company_name"), "推荐单位")
        pos = _safe(c.get("position"), "推荐岗位")
        score = c.get("match_score", "")
        ctype = _safe(c.get("cooperation_type"), "候选单位")
        salary = _safe(c.get("salary_treatment"), "以最新招聘信息为准")
        reason = _safe(c.get("match_reason"))
        core = c.get("core_functions") or []
        if isinstance(core, str):
            core = [core]

        _para(doc, f"{name}（{ctype}，匹配度 {score}）", bold=True)
        _para(doc, f"推荐岗位：{pos}", indent=0.5)
        _para(doc, f"薪资待遇：{salary}", indent=0.5)
        if core:
            _para(doc, "核心职能：" + "；".join(str(x) for x in core[:4]), indent=0.5)
        _para(doc, f"匹配原因：{reason}", indent=0.5, size=10, color="475569")
        doc.add_paragraph()


def _section_disclaimer(doc: Document):
    doc.add_page_break()
    _heading(doc, "免责声明", level=2)
    p = doc.add_paragraph(DISCLAIMER)
    p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


# ── public API ─────────────────────────────────────────────────────────────

def build_word_report(result: dict) -> io.BytesIO:
    """Build a Word report from a run_recommendation() result dict.

    Accepts both:
    - New structured student format (has gpa_by_semester) passed directly
    - Old run_recommendation() result dict with profile/context keys

    Args:
        result: dict with keys: profile, context, advice,
                algorithm_analysis, intent_analysis.
                OR a structured student dict with gpa_by_semester.

    Returns:
        BytesIO containing the .docx file.
    """
    # Detect if result is a structured student dict (from Excel import)
    if "gpa_by_semester" in result:
        profile_for_report = result  # pass the full structured student
        scores = result.get("scores") or {}
        path_scores = {}
        seniors = []
        companies = []
        algo = {}
        intent = {}
        advice = ""
    else:
        profile_for_report = result.get("profile") or {}
        context = result.get("context") or {}
        advice = result.get("advice") or ""
        algo = result.get("algorithm_analysis") or context.get("algorithm_analysis") or {}
        intent = result.get("intent_analysis") or context.get("intent_analysis") or {}
        scores = (context.get("current_student") or {}).get("scores") or profile_for_report.get("scores") or {}
        path_scores = context.get("path_scores") or {}
        seniors = context.get("similar_seniors") or []
        companies = context.get("recommended_companies") or []

    generated_at = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(11)

    _section_cover(doc, generated_at)
    _section_profile(doc, profile_for_report)
    _section_six_dim(doc, scores)
    if path_scores:
        _section_path_scores(doc, path_scores)
        _section_main_path(doc, path_scores, algo)
    _section_rationale(doc, algo, intent)
    _section_planning(doc, algo, advice)
    _section_seniors(doc, seniors)
    _section_companies(doc, companies)
    _section_disclaimer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def report_filename() -> str:
    return "student_growth_report_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx"
